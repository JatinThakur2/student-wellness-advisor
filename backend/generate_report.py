import json
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


PROJECT_NAME = "Student Wellness Advisor"
REPORT_FILENAME = "Student_Wellness_Advisor_Report.docx"


def load_test_results(path: str) -> dict:
    try:
        with open(path, "r") as f:
            return json.load(f)
    except FileNotFoundError:
        return {}


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    return h


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_bullets(doc: Document, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.size = Pt(11)


def build_report():
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{PROJECT_NAME} — Project Report")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
    subrun.italic = True
    subrun.font.size = Pt(11)

    doc.add_page_break()

    # Introduction
    add_heading(doc, "Introduction", 1)
    add_paragraph(
        doc,
        (
            "Student Wellness Advisor is a data-driven web application that estimates a student's "
            "well-being across three dimensions—Anxiety (GAD-7), Depression (PHQ-9), and Life Satisfaction—"
            "based on lifestyle, environment, and study-related inputs. The system combines a machine-learning "
            "regressor for scoring with an optional local LLM to generate supportive, actionable suggestions."
        ),
    )
    add_paragraph(
        doc,
        (
            "Objectives:\n"
            "• Collect relevant daily-life indicators from students in a simple UI.\n"
            "• Predict three wellness scores using a trained Random Forest model.\n"
            "• Provide a concise, empathetic summary and 3–5 suggestions via a local LLM (if available)."
        ),
    )

    # Data and preprocessing
    add_heading(doc, "Data and preprocessing", 1)
    add_paragraph(
        doc,
        (
            "Source: Student Wellness Survey (data).csv. Data is cleaned in backend/data_processing.py by:\n"
            "• Normalizing column names and trimming strings.\n"
            "• Converting categorical values (e.g., gender, roommate_situation) to numeric codes.\n"
            "• Extracting numeric values from mixed-format strings (e.g., '2–3 cups' → 2 or 3).\n"
            "• Casting scales to numeric and imputing remaining missing numeric values with column means."
        ),
    )

    # Model training
    add_heading(doc, "Model training", 1)
    features = [
        'age', 'gender', 'relationship_status', 'sleep_quality',
        'water_intake_litres', 'food_quality', 'meals_count',
        'caffeine_intake_cups', 'alcohol_consumption_drinks_week',
        'smoking_status', 'chronic_health_issues', 'medication_use',
        'physical_activity_minutes_day', 'sedentary_time_hours_day',
        'environment_noise', 'lighting_quality', 'study_space_quality',
        'roommate_situation', 'living_arrangement', 'commute_time_minutes_day',
        'study_hours_last_24h', 'study_stress',
        'upcoming_exam_assignment_stress', 'grades_estimate_gpa',
        'financial_stress', 'mood_last_24h', 'screen_time_hours_day',
        'social_media_time_hours_day', 'social_interaction_time_minutes_day'
    ]
    targets = ['anxiety_score_gad_7', 'depression_score_phq_9', 'life_satisfaction']

    add_paragraph(
        doc,
        (
            "A RandomForestRegressor (sklearn, 100 estimators, random_state=42) is trained on an 80/20 "
            "train/test split. The model predicts three continuous targets: Anxiety (GAD-7), Depression "
            "(PHQ-9), and Life Satisfaction. The trained model is persisted to backend/wellness_model.joblib."
        ),
    )
    add_paragraph(doc, f"Features used ({len(features)}): " + ", ".join(features))
    add_paragraph(doc, f"Targets: " + ", ".join(targets))

    # System architecture & workflow
    add_heading(doc, "System architecture and workflow", 1)
    add_paragraph(
        doc,
        (
            "The system follows a lightweight client–server architecture:\n"
            "1. Frontend (frontend/index.html) presents a form collecting 29 numeric/categorical inputs.\n"
            "2. Backend (Flask app in backend/app.py) exposes POST /predict. It:\n"
            "   • Loads the trained model (backend/wellness_model.joblib).\n"
            "   • Converts request JSON to a pandas DataFrame and runs model.predict().\n"
            "   • Optionally calls a local LLM via http://localhost:11434/api/generate to create a summary "
            "(model: 'deepseek-r1:1.5b') if an Ollama server is running.\n"
            "3. Response includes the three scores and an optional supportive summary."
        ),
    )

    # API contract
    add_heading(doc, "API contract (POST /predict)", 2)
    add_paragraph(
        doc,
        (
            "Request: JSON object with the 29 feature keys listed above (numbers only).\n"
            "Response: { anxiety_score: float, depression_score: float, life_satisfaction: float, summary: string }."
        ),
    )

    # How to run
    add_heading(doc, "How to run", 1)
    add_paragraph(
        doc,
        (
            "1) (Optional) Clean and prepare data: run backend/data_processing.py to generate backend/cleaned_data.csv.\n"
            "2) Train the model: python backend/train_model.py (produces wellness_model.joblib and backend/test_results.json).\n"
            "3) Start the backend: python backend/app.py (Flask on http://127.0.0.1:5000).\n"
            "4) Open frontend/index.html in a browser and submit the form.\n"
            "5) (Optional) For summaries, run a local LLM server compatible with Ollama API at http://localhost:11434."
        ),
    )

    # Testing and results
    add_heading(doc, "Testing and results", 1)
    results = load_test_results("backend/test_results.json")
    if results:
        per_target = results.get("per_target", {})
        summary = results.get("summary", {})

        # Per-target metrics table
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Target"
        hdr_cells[1].text = "R²"
        hdr_cells[2].text = "MAE"
        hdr_cells[3].text = "MSE"

        for tgt, vals in per_target.items():
            row_cells = table.add_row().cells
            row_cells[0].text = tgt
            row_cells[1].text = f"{vals.get('r2', float('nan')):.3f}"
            row_cells[2].text = f"{vals.get('mae', float('nan')):.3f}"
            row_cells[3].text = f"{vals.get('mse', float('nan')):.3f}"

        doc.add_paragraph("")
        add_paragraph(
            doc,
            (
                "Summary (uniform average across targets): "
                f"R²={summary.get('r2_uniform_average', float('nan')):.3f}, "
                f"MAE={summary.get('mae_uniform_average', float('nan')):.3f}, "
                f"MSE={summary.get('mse_uniform_average', float('nan')):.3f}"
            ),
        )
        add_paragraph(
            doc,
            (
                "Interpretation: R² values around 0.34–0.49 indicate moderate predictive power on this dataset. "
                "MAE and MSE provide absolute error magnitudes for each target; lower is better."
            ),
        )
    else:
        add_paragraph(
            doc,
            (
                "No test results file found (backend/test_results.json). Train the model to populate this section."
            ),
        )

    # Limitations & next steps
    add_heading(doc, "Limitations and future work", 1)
    add_bullets(
        doc,
        [
            "Survey size and self-reported noise may cap accuracy; collect more diverse data.",
            "Try models with calibrated uncertainty (e.g., Gradient Boosting, NGBoost) and hyperparameter tuning.",
            "Add input validation and schema checks on /predict.",
            "Enhance UI/UX and mobile responsiveness; add visualization of score distributions.",
            "Secure and cache LLM calls; allow model selection and prompt customization.",
        ],
    )

    # Appendix
    add_heading(doc, "Appendix: Key files", 1)
    add_bullets(
        doc,
        [
            "backend/data_processing.py — cleans raw survey data and exports cleaned_data.csv",
            "backend/train_model.py — trains RandomForestRegressor and saves wellness_model.joblib",
            "backend/app.py — Flask API with POST /predict and optional LLM summary",
            "backend/test_results.json — evaluation metrics on the test split",
            "frontend/index.html — simple HTML form and fetch logic",
        ],
    )

    doc.save(REPORT_FILENAME)
    return REPORT_FILENAME


if __name__ == "__main__":
    path = build_report()
    print(f"Report generated: {path}")
